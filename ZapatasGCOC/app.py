import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import io
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from zapatas_GCOC_1 import comprobacion_hundimiento

# --- CONFIGURACIÓN DE LA PÁGINA ---
st.set_page_config(page_title="Cimentaciones GCOC", page_icon="🏗️", layout="wide")

FS_MAP = {
    "Persistente (FS = 3.00)": 3.00,
    "Transitoria (FS = 2.50)": 2.50,
    "Accidental (FS = 2.00)": 2.00
}

# --- FUNCIÓN GENERADORA DEL REPORTE WORD ---
def generar_memoria_word(df_resultados, tipo, situacion, fs_obj):
    """Genera un documento Word en memoria con los resultados del análisis."""
    doc = Document()
    
    titulo = doc.add_heading('Memoria de Cálculo de Cimentación', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('1. Datos de Partida', level=1)
    doc.add_paragraph(f"Tipo de Cimentación analizada: Zapata {tipo}")
    doc.add_paragraph(f"Situación de Proyecto: {situacion} (Coeficiente de seguridad exigido: {fs_obj})")
    
    doc.add_heading('2. Resultados del Análisis Paramétrico', level=1)
    doc.add_paragraph(
        "A continuación se presenta el resumen de las iteraciones calculadas según "
        "la metodología de Brinch-Hansen modificada para la GCOC."
    )
    
    df_corte = df_resultados.head(30) # Limitamos a 30 combinaciones para no saturar el Word
    tabla = doc.add_table(rows=1, cols=len(df_corte.columns))
    tabla.style = 'Table Grid'
    
    hdr_cells = tabla.rows[0].cells
    for i, columna in enumerate(df_corte.columns):
        hdr_cells[i].text = str(columna)
        
    for _, fila in df_corte.iterrows():
        row_cells = tabla.add_row().cells
        for i, valor in enumerate(fila):
            row_cells[i].text = str(valor)
            
    if len(df_resultados) > 30:
        doc.add_paragraph(f"\n*Nota: Se muestran las primeras 30 combinaciones de un total de {len(df_resultados)} evaluadas en el cálculo.*")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- BARRA LATERAL ---
with st.sidebar:
    st.header("⚙️ Configuración Global")
    modo_operacion = st.radio("Modo de Operación:", ["A: Pre-dimensionamiento", "B: Verificación Estructural"])
    tipo_cimentacion = st.radio("Tipo de Geometría:", ["Rectangular", "Corrida", "Circular"])
    
    st.divider()
    situacion = st.selectbox("Situación de Proyecto:", options=list(FS_MAP.keys()))
    fs_objetivo = FS_MAP[situacion]
    drenaje = st.radio("Condición del terreno:", ["Largo Plazo (Drenado)", "Corto Plazo (No Drenado)"])

st.title(f"🏗️ Análisis de Zapata {tipo_cimentacion}")

# --- ENTRADA DE DATOS ---
col1, col2, col3 = st.columns(3)

with col1:
    with st.expander("🌍 Parámetros del Terreno", expanded=True):
        if drenaje == "Largo Plazo (Drenado)":
            c = st.number_input("Cohesión, c (kPa)", min_value=0.0, value=10.0)
            phi = st.number_input("Ángulo rozamiento, $\phi$ (°)", min_value=0.0, value=30.0)
        else:
            c = st.number_input("Resistencia al corte, $s_u$ (kPa)", min_value=0.0, value=50.0)
            phi = 0.0
            st.info("A corto plazo se asume $\phi = 0$")
            
        gamma_ap = st.number_input("Peso aparente, $\gamma_{ap}$ (kN/m³)", value=20.0)
        gamma_sub = st.number_input("Peso sumergido, $\gamma'$ (kN/m³)", value=10.0)
        gamma_sat = gamma_sub + 9.81

with col2:
    with st.expander("📏 Geometría Iterativa y NF", expanded=True):
        D = st.number_input("Profundidad apoyo, D (m)", min_value=0.0, value=1.5, step=0.1)
        hw = st.number_input("Nivel Freático bajo base, $h_w$ (m)", min_value=0.0, value=5.0)
        D_w = D + hw
        
        etiqueta_B = "Diámetro D (m)" if tipo_cimentacion == "Circular" else "Ancho B (m)"
        st.markdown(f"**Rango de {etiqueta_B}:**")
        B_min, B_max, B_inc = st.columns(3)
        B_min = B_min.number_input(f"Mínimo", min_value=0.5, value=1.0, step=0.5, key="b1")
        B_max = B_max.number_input(f"Máximo", min_value=1.0, value=4.0, step=0.5, key="b2")
        B_inc = B_inc.number_input(f"Inc.", min_value=0.1, value=0.1, step=0.1, key="b3")
        
        if tipo_cimentacion == "Rectangular":
            st.markdown("**Rango de Longitud L (m):**")
            L_min, L_max, L_inc = st.columns(3)
            L_min = L_min.number_input("L Mín", min_value=0.5, value=1.0, step=0.5)
            L_max = L_max.number_input("L Máx", min_value=1.0, value=5.0, step=0.5)
            L_inc = L_inc.number_input("L Inc", min_value=0.1, value=0.5, step=0.1)
        else:
            L_min, L_max, L_inc = 1.0, 1.0, 1.0  # Dummy para el bucle unidimensional

with col3:
    with st.expander("⚖️ Cargas de la Estructura", expanded=True):
        if modo_operacion.startswith("B"):
            if tipo_cimentacion == "Corrida":
                st.caption("Nota: Cargas expresadas **por metro lineal**.")
            
            V = st.number_input("Vertical, V", min_value=0.1, value=1000.0)
            H = st.number_input("Horizontal, H", value=0.0)
            M_B = st.number_input(f"Momento flector en {etiqueta_B[0]}, $M_B$", value=0.0)
            
            if tipo_cimentacion == "Rectangular":
                M_L = st.number_input("Momento flector en L, $M_L$", value=0.0)
            else:
                M_L = 0.0
        else:
            V, H, M_B, M_L = 1.0, 0.0, 0.0, 0.0
            st.info("Modo A: Generación de la Carta de Tensiones paramétrica sin cargas.")

# --- MOTOR ITERATIVO ---
resultados = []

for B in np.arange(B_min, B_max + B_inc, B_inc):
    for L in np.arange(L_min, L_max + L_inc, L_inc):
        if tipo_cimentacion == "Rectangular" and L < B:
            continue
            
        e_B = abs(M_B / V) if V > 0 else 0.0
        e_L = abs(M_L / V) if V > 0 else 0.0
        
        B_star = B - 2 * e_B
        if tipo_cimentacion == "Rectangular":
            L_star = L - 2 * e_L
        else:
            L_star = 1.0 
            
        if B_star <= 0 or (tipo_cimentacion == "Rectangular" and L_star <= 0):
            continue 

        res = comprobacion_hundimiento(
            V=V, H=H, c=c, phi_deg=phi, gamma_ap=gamma_ap, 
            gamma_sat=gamma_sat, D_w=D_w, D=D, 
            B_star=B_star, L_star=L_star, F_h_exigido=fs_objetivo,
            tipo_cimentacion=tipo_cimentacion
        )
        
        fila = {etiqueta_B: round(B, 3)}
        if tipo_cimentacion == "Rectangular":
            fila["L (m)"] = round(L, 3)
            
        if modo_operacion.startswith("B"):
            fila.update({
                "Área Ef. (m²)": round(res.area_efectiva, 3), 
                "p_actuante (kPa)": round(res.p_v_actuante, 1), 
                "p_hundimiento (kPa)": round(res.p_vh, 1), 
                "FS": round(res.F_real, 2), 
                "Cumple": "✅ Sí" if res.cumple_normativa else "❌ No"
            })
        else:
            fila.update({
                "p_hundimiento (kPa)": round(res.p_vh, 1), 
                "p_admisible (kPa)": round(res.p_v_adm, 1)
            })
        resultados.append(fila)

df_res = pd.DataFrame(resultados)

# --- VISUALIZACIÓN Y REPORTES ---
st.header("📊 Resultados del Análisis")
if df_res.empty:
    st.error("⚠️ Vuelco total de la cimentación para los rangos dados.")
else:
    tab_grafica, tab_datos, tab_exportar = st.tabs(["📈 Gráfica Paramétrica", "📋 Tabla de Datos", "📥 Exportar Memoria"])
    
    with tab_grafica:
        if tipo_cimentacion == "Rectangular":
            if modo_operacion.startswith("A"):
                matriz_z = df_res.pivot(index="L (m)", columns=etiqueta_B, values="p_admisible (kPa)")
                fig = go.Figure(data=go.Heatmap(z=matriz_z.values, x=matriz_z.columns, y=matriz_z.index, colorscale="Viridis"))
                fig.update_layout(title="Mapa de Calor: Tensión Admisible (kPa)", xaxis_title="B (m)", yaxis_title="L (m)")
                st.plotly_chart(fig, use_container_width=True)
            else:
                fig = px.scatter(df_res, x=etiqueta_B, y="L (m)", color="FS", color_continuous_scale="RdYlGn", range_color=[fs_objetivo-1, fs_objetivo+2])
                st.plotly_chart(fig, use_container_width=True)
        else: 
            y_col = "p_admisible (kPa)" if modo_operacion.startswith("A") else "FS"
            fig = px.line(df_res, x=etiqueta_B, y=y_col, markers=True, title=f"Evolución de {y_col} frente a la dimensión principal")
            
            if modo_operacion.startswith("B"):
                fig.add_hline(y=fs_objetivo, line_dash="dash", line_color="red", annotation_text=f"FS Mínimo = {fs_objetivo}")
                
            fig.update_layout(xaxis_title=etiqueta_B, yaxis_title=y_col)
            st.plotly_chart(fig, use_container_width=True)

    with tab_datos:
        st.dataframe(df_res, use_container_width=True, hide_index=True)

    with tab_exportar:
        st.subheader("Generación de Entregables")
        
        col_csv, col_word = st.columns(2)
        
        with col_csv:
            st.markdown("**Datos crudos (CSV)**")
            st.download_button(
                label="💾 Descargar CSV", 
                data=df_res.to_csv(index=False).encode('utf-8'), 
                file_name='analisis_zapatas.csv', 
                mime='text/csv'
            )
            
        with col_word:
            st.markdown("**Memoria de Cálculo (Word)**")
            
            if st.button("⚙️ Generar Memoria Word"):
                with st.spinner("Redactando anejo de cálculo..."):
                    buffer_word = generar_memoria_word(
                        df_res, tipo_cimentacion, situacion, fs_objetivo
                    )
                    st.session_state['archivo_word'] = buffer_word
                st.success("¡Memoria generada correctamente!")
            
            if 'archivo_word' in st.session_state:
                st.download_button(
                    label="📄 Descargar Memoria.docx",
                    data=st.session_state['archivo_word'],
                    file_name=f"Memoria_Zapatas_{tipo_cimentacion}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )