import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import io
from fpdf import FPDF
from docx import Document

from motor_calculo import Terreno, Cargas, Cimentacion, CalculadoraCapacidad

# --- FUNCIONES DE GENERACIÓN DE INFORMES ---
def generar_pdf(resultado: dict) -> io.BytesIO:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    det = resultado['Detalles']
    
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Memoria Justificativa de Cimentación (CTE-DB-SE-C)", ln=True, align='C')
    pdf.ln(5)
    
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(200, 8, txt="1. Geometría y Resultado Final", ln=True)
    pdf.set_font("Arial", size=10)
    pdf.cell(200, 6, txt=f"Ancho / Diámetro (B): {resultado['B']} m | Longitud (L): {resultado['L']} m", ln=True)
    pdf.cell(200, 6, txt=f"Ancho Efectivo (B*): {resultado['B*']:.2f} m | Área Efectiva: {resultado['Área Ef.']:.2f} m2", ln=True)
    pdf.cell(200, 6, txt=f"Presión de Trabajo: {resultado['q_trabajo']:.2f} kPa", ln=True)
    pdf.cell(200, 6, txt=f"Presión Admisible: {resultado['q_adm']:.2f} kPa", ln=True)
    estado = "CUMPLE" if resultado['Cumple'] else "NO CUMPLE"
    pdf.cell(200, 6, txt=f"Verificación ELU de Hundimiento: {estado}", ln=True)
    pdf.ln(4)
    
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(200, 8, txt="2. Coeficientes del Método Analítico", ln=True)
    pdf.set_font("Arial", size=10)
    pdf.cell(200, 6, txt=f"Capacidad de Carga: Nq={det['Nq']:.2f}, Nc={det['Nc']:.2f}, Ngamma={det['Ng']:.2f}", ln=True)
    pdf.cell(200, 6, txt=f"Factores de Forma: sq={det['sq']:.2f}, sc={det['sc']:.2f}, sgamma={det['sg']:.2f}", ln=True)
    pdf.cell(200, 6, txt=f"Factores de Profundidad: dq={det['dq']:.2f}, dc={det['dc']:.2f}, dgamma={det['dg']:.2f}", ln=True)
    pdf.cell(200, 6, txt=f"Factores de Inclinación: iq={det['iq']:.2f}, ic={det['ic']:.2f}, igamma={det['ig']:.2f}", ln=True)
    pdf.cell(200, 6, txt=f"Factores de Talud: tq={det['tq']:.2f}, tc={det['tc']:.2f}, tgamma={det['tg']:.2f}", ln=True)
    
    pdf_buffer = io.BytesIO()
    pdf_buffer.write(pdf.output())
    pdf_buffer.seek(0)
    return pdf_buffer

def generar_docx(resultado: dict) -> io.BytesIO:
    doc = Document()
    doc.add_heading('Memoria Justificativa de Cimentación (CTE-DB-SE-C)', 0)
    det = resultado['Detalles']
    
    doc.add_heading('1. Geometría y Resultado Final', level=1)
    doc.add_paragraph(f"Ancho / Diámetro (B): {resultado['B']} m | Longitud (L): {resultado['L']} m")
    doc.add_paragraph(f"Ancho Efectivo (B*): {resultado['B*']:.2f} m | Área Efectiva: {resultado['Área Ef.']:.2f} m2")
    doc.add_paragraph(f"Presión de Trabajo: {resultado['q_trabajo']:.2f} kPa")
    doc.add_paragraph(f"Presión Admisible: {resultado['q_adm']:.2f} kPa")
    estado = "CUMPLE" if resultado['Cumple'] else "NO CUMPLE"
    doc.add_paragraph(f"Verificación ELU de Hundimiento: {estado}")
    
    doc.add_heading('2. Coeficientes del Método Analítico', level=1)
    doc.add_paragraph(f"Capacidad de Carga: Nq={det['Nq']:.2f}, Nc={det['Nc']:.2f}, Ngamma={det['Ng']:.2f}")
    doc.add_paragraph(f"Factores de Forma: sq={det['sq']:.2f}, sc={det['sc']:.2f}, sgamma={det['sg']:.2f}")
    doc.add_paragraph(f"Factores de Profundidad: dq={det['dq']:.2f}, dc={det['dc']:.2f}, dgamma={det['dg']:.2f}")
    doc.add_paragraph(f"Factores de Inclinación: iq={det['iq']:.2f}, ic={det['ic']:.2f}, igamma={det['ig']:.2f}")
    doc.add_paragraph(f"Factores de Talud: tq={det['tq']:.2f}, tc={det['tc']:.2f}, tgamma={det['tg']:.2f}")
    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

# --- INTERFAZ DE STREAMLIT ---
st.set_page_config(page_title="CTE DB SE-C: Hundimiento", layout="wide")
st.title("🏗️ Cálculo Analítico de Hundimiento (CTE DB SE-C)")

with st.sidebar:
    st.header("1. Terreno y Agua")
    es_drenado = st.checkbox("Análisis Drenado", value=True)
    c_0 = st.number_input("Cohesión [kPa]", min_value=0.0, value=10.0)
    m_grad = st.number_input("Gradiente m [kPa/m]", min_value=0.0, value=0.0) if not es_drenado else 0.0
    phi_k = st.number_input("Áng. Rozamiento φ' [°]", min_value=0.0, value=30.0) if es_drenado else 0.0
    gamma_ap = st.number_input("γ aparente [kN/m³]", min_value=0.1, value=18.0)
    gamma_sat = st.number_input("γ saturado [kN/m³]", min_value=0.1, value=20.0)
    
    st.markdown("---")
    z_w = st.number_input("Profundidad N.F. [m]", min_value=0.0, value=5.0)
    i_v = st.number_input("Gradiente ascendente i_v", min_value=0.0, value=0.0)
    ignorar_coef_prof = st.checkbox("Anular resistencia del terreno superior (Coef. d_i = 1)", value=False)

    st.header("2. Cimentación")
    tipo_cim = st.selectbox("Tipo de Zapata", ["rectangular", "corrida", "circular"])
    D = st.number_input("Cota de apoyo D [m]", min_value=0.0, value=1.0)
    q_ext = st.number_input("Sobrecarga permanente q_ext [kPa]", min_value=0.0, value=0.0)
    beta = st.number_input("Inclinación talud β [°]", min_value=0.0, max_value=45.0, value=0.0)
    c_base = st.number_input("Cohesión base-terreno c [kPa]", min_value=0.0, value=0.0)

    st.header("3. Acciones")
    V = st.number_input("Axil V [kN]", min_value=0.1, value=1000.0)
    H = st.number_input("Cortante H [kN]", min_value=0.0, value=0.0)
    M_B = st.number_input("Momento M_B [mkN]", value=0.0)
    M_L = st.number_input("Momento M_L [mkN]", value=0.0) if tipo_cim in ['rectangular', 'circular'] else 0.0

st.write("### Barrido Geométrico")
c1, c2, c3 = st.columns(3)
label_B = "Diámetro mín [m]" if tipo_cim == 'circular' else "Ancho B mín [m]"
B_min = c1.number_input(label_B, value=1.0)
B_max = c2.number_input(label_B.replace("mín", "máx"), value=4.0)
paso = c3.number_input("Paso [m]", value=0.25)

if tipo_cim == 'rectangular':
    L_min = c1.number_input("Longitud L mín [m]", value=1.0)
    L_max = c2.number_input("Longitud L máx [m]", value=4.0)

if st.button("🚀 Calcular Matriz", type="primary"):
    terreno = Terreno(c_0, m_grad, phi_k, gamma_ap, gamma_sat, z_w, es_drenado, i_v, ignorar_coef_prof)
    cargas = Cargas(V, H, M_B, M_L)
    cimentacion = Cimentacion(D, beta, tipo_cim, c_base, q_ext)
    calculadora = CalculadoraCapacidad(terreno, cimentacion, cargas)

    resultados = []
    rango_B = np.arange(B_min, B_max + paso/2, paso)
    
    if tipo_cim == 'rectangular':
        rango_L = np.arange(L_min, L_max + paso/2, paso)
        for b in rango_B:
            for l in rango_L:
                if b <= l: 
                    res = calculadora.calcular(b, l)
                    if res: resultados.append(res)
    else:
        for b in rango_B:
            res = calculadora.calcular(b)
            if res: resultados.append(res)

    if resultados:
        df = pd.DataFrame(resultados)
        tab1, tab2, tab3, tab4 = st.tabs(["📊 Resultados", "📈 Gráfico", "🔎 Desglose", "📄 Informes"])
        
        with tab1:
            cols = ["B", "B*", "Área Ef.", "q_trabajo", "q_h", "q_adm", "Cumple"]
            if tipo_cim == 'rectangular': cols.insert(1, "L")
            st.dataframe(df[cols].style.map(lambda x: 'background-color: #d4edda' if x else 'background-color: #f8d7da', subset=['Cumple']), use_container_width=True)

        with tab2:
            fig, ax = plt.subplots(figsize=(10, 4))
            if tipo_cim == 'rectangular':
                for l in np.unique(df['L']):
                    subset = df[df['L'] == l]
                    ax.plot(subset['B'], subset['q_adm'], marker='.', label=f'q_adm (L={l}m)')
            else:
                ax.plot(df['B'], df['q_adm'], marker='o', label='Capacidad Admisible')
            ax.plot(df['B'], df['q_trabajo'], linestyle='--', color='red', label='Presión Actuante')
            ax.set_xlabel('Dimensión Principal (m)')
            ax.set_ylabel('Presión (kPa)')
            ax.grid(True)
            ax.legend()
            st.pyplot(fig)

        with tab3:
            etiquetas = [f"B={r['B']}m, L={r['L']}m" if tipo_cim == 'rectangular' else (f"D={r['B']}m" if tipo_cim == 'circular' else f"B={r['B']}m (Corrida)") for r in resultados]
            seleccion = st.selectbox("Seleccione zapata:", etiquetas)
            det = resultados[etiquetas.index(seleccion)]['Detalles']
            
            c_a, c_b, c_c = st.columns(3)
            c_a.write(f"**Nq:** {det['Nq']:.2f} | **Nc:** {det['Nc']:.2f} | **Nγ:** {det['Ng']:.2f}")
            c_b.write(f"**sq:** {det['sq']:.2f} | **sc:** {det['sc']:.2f} | **sγ:** {det['sg']:.2f}")
            c_c.write(f"**dq:** {det['dq']:.2f} | **dc:** {det['dc']:.2f} | **dγ:** {det['dg']:.2f}")
            c_a.write(f"**iq:** {det['iq']:.2f} | **ic:** {det['ic']:.2f} | **iγ:** {det['ig']:.2f}")
            c_b.write(f"**tq:** {det['tq']:.2f} | **tc:** {det['tc']:.2f} | **tγ:** {det['tg']:.2f}")
            c_c.write(f"**q0k:** {det['q0k']:.1f} kPa | **γ_k:** {det['gamma_k']:.1f} kN/m³")

        with tab4:
            st.info("Descargar memoria justificativa (ELU Hundimiento):")
            seleccion_inf = st.selectbox("Geometría:", etiquetas, key="inf")
            resultado_elegido = resultados[etiquetas.index(seleccion_inf)]
            
            c_pdf, c_docx = st.columns(2)
            c_pdf.download_button("📥 PDF", data=generar_pdf(resultado_elegido), file_name="Memoria.pdf", mime="application/pdf")
            c_docx.download_button("📥 Word", data=generar_docx(resultado_elegido), file_name="Memoria.docx", mime="application/vnd.openxmlformats")
    else:
        st.error("No hay soluciones válidas. Comprueba las excentricidades.")
